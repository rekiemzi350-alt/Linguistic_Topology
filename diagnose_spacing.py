import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

path = os.path.expanduser("~/coffee/12-11-2025 Conversation with Gemini on AI and Computer Games.docx")

def get_alignment_name(alignment):
    if alignment == WD_ALIGN_PARAGRAPH.LEFT: return "LEFT"
    if alignment == WD_ALIGN_PARAGRAPH.CENTER: return "CENTER"
    if alignment == WD_ALIGN_PARAGRAPH.RIGHT: return "RIGHT"
    if alignment == WD_ALIGN_PARAGRAPH.JUSTIFY: return "JUSTIFY"
    return str(alignment)

try:
    doc = Document(path)
    print(f"Total Paragraphs: {len(doc.paragraphs)}")
    
    print("\n--- First 5 Paragraphs ---")
    for i in range(min(5, len(doc.paragraphs))):
        p = doc.paragraphs[i]
        print(f"[{i}] Align: {get_alignment_name(p.alignment)}")
        print(f"    Text: {p.text[:50]}...")

    print("\n--- Paragraphs 50-55 ---")
    for i in range(50, min(56, len(doc.paragraphs))):
        p = doc.paragraphs[i]
        print(f"[{i}] Align: {get_alignment_name(p.alignment)}")
        print(f"    Text: {p.text[:50]}...")

except Exception as e:
    print(f"Error: {e}")
