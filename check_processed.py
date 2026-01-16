import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

path = "/storage/emulated/0/Books/AI_Conversations/Gemini on AI.docx"

def get_alignment_name(alignment):
    if alignment == WD_ALIGN_PARAGRAPH.LEFT: return "LEFT"
    if alignment == WD_ALIGN_PARAGRAPH.CENTER: return "CENTER"
    return str(alignment)

try:
    doc = Document(path)
    print(f"Checking: {path}")
    for i in range(min(5, len(doc.paragraphs))):
        p = doc.paragraphs[i]
        print(f"[{i}] Align: {get_alignment_name(p.alignment)} | Text: {p.text[:30]}")
except Exception as e:
    print(e)
