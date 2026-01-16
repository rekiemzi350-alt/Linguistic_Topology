import os
import glob
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Paths
REF_PATH = os.path.expanduser("~/coffee/12-11-2025 Conversation with Gemini on AI and Computer Games.docx")
TARGET_DIR = "/storage/emulated/0/Books/AI_Conversations/"

def get_style_from_ref(path):
    try:
        doc = Document(path)
        font_name = None
        font_size = None

        # Check the 'Normal' style first as a baseline
        style = doc.styles['Normal']
        if style.font.name:
            font_name = style.font.name
        if style.font.size:
            font_size = style.font.size

        # Check first non-empty paragraph for direct formatting overrides
        for p in doc.paragraphs:
            if p.text.strip():
                for run in p.runs:
                    if run.font.name:
                        font_name = run.font.name
                    if run.font.size:
                        font_size = run.font.size
                break # Stop after checking the first paragraph
        
        return font_name, font_size
    except Exception as e:
        print(f"Warning: Could not read reference style completely: {e}")
        return None, None

def process_files():
    # Get reference style
    ref_font, ref_size = get_style_from_ref(REF_PATH)
    
    # Defaults if detection fails
    if not ref_font: ref_font = "Times New Roman"
    if not ref_size: ref_size = Pt(12)
    
    print(f"Using Font: {ref_font}, Size: {ref_size}")

    docx_files = glob.glob(os.path.join(TARGET_DIR, "*.docx"))

    for file_path in docx_files:
        filename = os.path.basename(file_path)
        # Skip the reference file if it happens to be in the folder (though it's in coffee currently)
        if os.path.abspath(file_path) == os.path.abspath(REF_PATH):
            continue

        try:
            print(f"Processing: {filename}")
            original_doc = Document(file_path)
            new_doc = Document()

            # 1. Title (Filename, Centered)
            title_text = os.path.splitext(filename)[0]
            title_para = new_doc.add_paragraph(title_text)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in title_para.runs:
                run.font.name = ref_font
                run.font.size = ref_size

            # 2. Body Paragraphs
            for p in original_doc.paragraphs:
                text = p.text.strip()
                if not text:
                    continue # Skip empty lines

                # Format Prefixes
                # User wants "Erik >" and "Gemini *"
                if text.startswith(">"):
                    if not text.startswith("Erik >") and not text.startswith("Erik>"):
                         text = "Erik " + text
                elif text.startswith("*"):
                    if not text.startswith("Gemini *") and not text.startswith("Gemini*"):
                        text = "Gemini " + text
                
                # Add Paragraph
                new_para = new_doc.add_paragraph(text)
                new_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Apply Font Style
                for run in new_para.runs:
                    run.font.name = ref_font
                    run.font.size = ref_size

            new_doc.save(file_path)
            print(f"  - Done")

        except Exception as e:
            print(f"  - Failed: {e}")

if __name__ == "__main__":
    process_files()
