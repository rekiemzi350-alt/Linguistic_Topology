import os
import glob
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Paths
REF_FILE = "/data/data/com.termux/files/home/coffee/12-11-2025 Conversation with Gemini on AI and Computer Games.docx"
TARGET_DIR = "/storage/emulated/0/Books/AI_Conversations/"

def get_style_from_ref(path):
    try:
        doc = Document(path)
        font_name = None
        font_size = None

        # Try to find style from the first few paragraphs
        for p in doc.paragraphs:
            if p.text.strip():
                for run in p.runs:
                    if run.font.name:
                        font_name = run.font.name
                    if run.font.size:
                        font_size = run.font.size
                if font_name and font_size:
                    break
        
        return font_name, font_size
    except Exception as e:
        print(f"Warning: Could not read reference style: {e}")
        return None, None

def process_file(file_path, ref_font, ref_size, is_reference_file=False):
    print(f"Processing: {os.path.basename(file_path)}")
    try:
        original_doc = Document(file_path)
        new_doc = Document()

        # 1. Title (Filename, Centered)
        # For the reference file, we might want to keep the original title if it exists, 
        # but the user said "format the rest of the file in the same fashion", implies uniformity.
        # I'll stick to the "Filename as Title" rule for consistency unless it looks weird.
        filename = os.path.basename(file_path)
        title_text = os.path.splitext(filename)[0]
        title_para = new_doc.add_paragraph(title_text)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title_para.runs:
            run.font.name = ref_font
            run.font.size = ref_size
            run.bold = True # Optional: make title bold? The user didn't ask, but it's standard. I'll leave it plain as requested "centered... bodies to left".

        # 2. Body Paragraphs
        for p in original_doc.paragraphs:
            raw_text = p.text
            if not raw_text.strip():
                continue 

            # Fix Spacing: Collapse multiple spaces to one
            clean_text = re.sub(r'\s+', ' ', raw_text).strip()

            # Fix Prefixes
            if clean_text.startswith(">"):
                if not clean_text.startswith("Erik >") and not clean_text.startswith("Erik>"):
                        clean_text = "Erik " + clean_text
            elif clean_text.startswith("*"):
                if not clean_text.startswith("Gemini *") and not clean_text.startswith("Gemini*"):
                    clean_text = "Gemini " + clean_text
            
            # Add Paragraph
            new_para = new_doc.add_paragraph(clean_text)
            new_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Apply Font Style
            for run in new_para.runs:
                run.font.name = ref_font
                run.font.size = ref_size

        new_doc.save(file_path)
        print(f"  - Done")

    except Exception as e:
        print(f"  - Failed: {e}")

def main():
    # Get style
    ref_font, ref_size = get_style_from_ref(REF_FILE)
    if not ref_font: ref_font = "Georgia" # Fallback
    if not ref_size: ref_size = Pt(12)    # Fallback
    
    print(f"Using Font: {ref_font}, Size: {ref_size}")

    # Process Reference File
    process_file(REF_FILE, ref_font, ref_size, is_reference_file=True)

    # Process Target Directory Files
    docx_files = glob.glob(os.path.join(TARGET_DIR, "*.docx"))
    for file_path in docx_files:
        # Avoid processing the reference file again if it happens to be in the list (it isn't, but safety first)
        if os.path.abspath(file_path) == os.path.abspath(REF_FILE):
            continue
        process_file(file_path, ref_font, ref_size)

if __name__ == "__main__":
    main()
